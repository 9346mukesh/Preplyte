"""Day 4 tests: grounded classification, verification pass, and LLM integration.

Tests:
- classify_requirements: PRESENT / PARTIAL / MISSING classifications (rule-based)
- verification_pass: evidence validation, INSUFFICIENT_EVIDENCE downgrade
- LLM classification: mock Groq response parsing, fallback on error
- End-to-end pipeline with real embeddings and classification
"""

from unittest.mock import MagicMock, patch

import pytest

from app.schemas import (
    AnalysisResult,
    Classification,
    JDRequirement,
    RequirementCategory,
    RetrievalResult,
    ResumeChunk,
    Section,
)
from app.services.analyzer import (
    PRESENT_THRESHOLD,
    _parse_llm_response,
    classify_requirements,
    verification_pass,
)


# ---------------------------------------------------------------------------
# classify_requirements tests (FR-05 / FR-06) — rule-based fallback
# ---------------------------------------------------------------------------


class TestClassifyRequirements:
    """FR-05: Classify each JD requirement with cited evidence."""

    def _make_chunks(self):
        return [
            ResumeChunk(
                chunk_id="skills:0",
                source_section=Section.SKILLS,
                raw_text="Python, FastAPI, PostgreSQL, Redis",
            ),
            ResumeChunk(
                chunk_id="exp:0",
                source_section=Section.EXPERIENCE,
                raw_text="Built REST APIs with FastAPI and PostgreSQL",
            ),
            ResumeChunk(
                chunk_id="edu:0",
                source_section=Section.EDUCATION,
                raw_text="B.Tech CSE, GITAM University",
            ),
        ]

    def test_present_when_high_similarity(self):
        """High score above PRESENT_THRESHOLD -> PRESENT with citation."""
        chunks = self._make_chunks()
        reqs = [JDRequirement(requirement_id="r1", requirement_text="Python experience")]
        retrieval = [
            RetrievalResult(
                requirement_id="r1",
                retrieved_chunk_ids=["skills:0"],
                similarity_scores=[0.82],
                threshold_pass=True,
            )
        ]
        results = classify_requirements(reqs, retrieval, chunks)
        assert len(results) == 1
        assert results[0].classification == Classification.PRESENT
        assert "Python" in results[0].evidence_citation
        assert results[0].confidence_note is not None

    def test_partial_when_moderate_similarity(self):
        """Moderate score between threshold and PRESENT_THRESHOLD -> PARTIAL."""
        chunks = self._make_chunks()
        reqs = [JDRequirement(requirement_id="r1", requirement_text="Docker containerization")]
        retrieval = [
            RetrievalResult(
                requirement_id="r1",
                retrieved_chunk_ids=["skills:0"],
                similarity_scores=[0.55],
                threshold_pass=True,
            )
        ]
        results = classify_requirements(reqs, retrieval, chunks)
        assert len(results) == 1
        assert results[0].classification == Classification.PARTIAL
        assert results[0].evidence_citation is not None

    def test_missing_when_threshold_fails(self):
        """Below threshold -> MISSING with no citation."""
        chunks = self._make_chunks()
        reqs = [JDRequirement(requirement_id="r1", requirement_text="Kubernetes orchestration")]
        retrieval = [
            RetrievalResult(
                requirement_id="r1",
                retrieved_chunk_ids=["edu:0"],
                similarity_scores=[0.20],
                threshold_pass=False,
            )
        ]
        results = classify_requirements(reqs, retrieval, chunks)
        assert len(results) == 1
        assert results[0].classification == Classification.MISSING
        assert results[0].evidence_citation is None

    def test_missing_when_no_retrieval(self):
        """No retrieval result -> MISSING."""
        chunks = self._make_chunks()
        reqs = [JDRequirement(requirement_id="r1", requirement_text="ML pipeline")]
        results = classify_requirements(reqs, [], chunks)
        assert len(results) == 1
        assert results[0].classification == Classification.MISSING

    def test_multiple_requirements(self):
        """Multiple requirements classified independently."""
        chunks = self._make_chunks()
        reqs = [
            JDRequirement(requirement_id="r1", requirement_text="Python"),
            JDRequirement(requirement_id="r2", requirement_text="Kubernetes"),
        ]
        retrieval = [
            RetrievalResult(
                requirement_id="r1",
                retrieved_chunk_ids=["skills:0"],
                similarity_scores=[0.80],
                threshold_pass=True,
            ),
            RetrievalResult(
                requirement_id="r2",
                retrieved_chunk_ids=[],
                similarity_scores=[],
                threshold_pass=False,
            ),
        ]
        results = classify_requirements(reqs, retrieval, chunks)
        assert results[0].classification == Classification.PRESENT
        assert results[1].classification == Classification.MISSING


# ---------------------------------------------------------------------------
# verification_pass tests (FR-06)
# ---------------------------------------------------------------------------


class TestVerificationPass:
    """FR-06: Re-check each claim against its cited evidence."""

    def test_present_with_valid_evidence(self):
        """Present claim with substantial evidence passes verification."""
        analysis = AnalysisResult(
            requirement_id="r1",
            classification=Classification.PRESENT,
            evidence_citation="Python, FastAPI, PostgreSQL, Redis — built REST APIs",
            confidence_note="Strong match (0.82)",
        )
        results = verification_pass([analysis])
        assert results[0].classification == Classification.PRESENT

    def test_downgrade_when_evidence_thin(self):
        """Evidence with <3 tokens -> INSUFFICIENT_EVIDENCE."""
        analysis = AnalysisResult(
            requirement_id="r1",
            classification=Classification.PRESENT,
            evidence_citation="N/A",
            confidence_note="Strong match",
        )
        results = verification_pass([analysis])
        assert results[0].classification == Classification.INSUFFICIENT_EVIDENCE
        assert "thin" in results[0].confidence_note.lower()

    def test_downgrade_when_no_citation(self):
        """Present claim with no citation -> INSUFFICIENT_EVIDENCE."""
        analysis = AnalysisResult(
            requirement_id="r1",
            classification=Classification.PARTIAL,
            evidence_citation=None,
            confidence_note="Partial match",
        )
        results = verification_pass([analysis])
        assert results[0].classification == Classification.INSUFFICIENT_EVIDENCE

    def test_missing_skipped(self):
        """Missing claims are not verified (already abstained)."""
        analysis = AnalysisResult(
            requirement_id="r1",
            classification=Classification.MISSING,
            evidence_citation=None,
            confidence_note="No relevant resume content found.",
        )
        results = verification_pass([analysis])
        assert results[0].classification == Classification.MISSING

    def test_insufficient_evidence_skipped(self):
        """INSUFFICIENT_EVIDENCE claims pass through unchanged."""
        analysis = AnalysisResult(
            requirement_id="r1",
            classification=Classification.INSUFFICIENT_EVIDENCE,
            confidence_note="Already abstained.",
        )
        results = verification_pass([analysis])
        assert results[0].classification == Classification.INSUFFICIENT_EVIDENCE


# ---------------------------------------------------------------------------
# LLM response parsing tests
# ---------------------------------------------------------------------------


class TestLLMResponseParsing:
    """Test JSON extraction from various LLM response formats."""

    def test_parse_clean_json(self):
        """Direct JSON parse should work."""
        response = '{"classification": "present", "evidence_citation": "Python, FastAPI", "confidence_note": "Strong match"}'
        parsed = _parse_llm_response(response)
        assert parsed is not None
        assert parsed["classification"] == "present"
        assert parsed["evidence_citation"] == "Python, FastAPI"

    def test_parse_markdown_fenced_json(self):
        """JSON inside markdown code fences should be extracted."""
        response = '```json\n{"classification": "partial", "evidence_citation": "Some evidence", "confidence_note": "Partial match"}\n```'
        parsed = _parse_llm_response(response)
        assert parsed is not None
        assert parsed["classification"] == "partial"

    def test_parse_with_surrounding_text(self):
        """JSON embedded in explanatory text should be found."""
        response = 'Based on the analysis, here is the result:\n{"classification": "missing", "evidence_citation": null, "confidence_note": "No match"}\nThis is my conclusion.'
        parsed = _parse_llm_response(response)
        assert parsed is not None
        assert parsed["classification"] == "missing"

    def test_parse_invalid_json(self):
        """Invalid JSON should return None."""
        parsed = _parse_llm_response("This is not JSON at all")
        assert parsed is None


# ---------------------------------------------------------------------------
# LLM classification integration tests
# ---------------------------------------------------------------------------


class TestLLMClassification:
    """Test LLM-based classification with mocked Groq API."""

    def test_llm_present_classification(self):
        """Mocked LLM returning 'present' should produce PRESENT result."""
        mock_response = MagicMock()
        mock_response.content = '{"classification": "present", "evidence_citation": "5 years of Python experience", "confidence_note": "Direct match"}'

        with patch("langchain_groq.ChatGroq") as MockChatGroq:
            mock_llm = MagicMock()
            mock_llm.invoke.return_value = mock_response
            MockChatGroq.return_value = mock_llm

            with patch("app.services.analyzer.get_settings") as mock_settings:
                mock_settings.return_value.groq_api_key = "test-key"
                mock_settings.return_value.llm_model = "llama-3.3-70b-versatile"

                chunks = [
                    ResumeChunk(chunk_id="s:0", source_section=Section.SKILLS, raw_text="Python, FastAPI")
                ]
                reqs = [JDRequirement(requirement_id="r1", requirement_text="Python experience")]
                retrieval = [
                    RetrievalResult(
                        requirement_id="r1",
                        retrieved_chunk_ids=["s:0"],
                        similarity_scores=[0.85],
                        threshold_pass=True,
                    )
                ]
                results = classify_requirements(reqs, retrieval, chunks)

                assert results[0].classification == Classification.PRESENT
                assert results[0].evidence_citation == "5 years of Python experience"

    def test_llm_missing_enforces_no_citation(self):
        """LLM returning 'present' with no citation -> INSUFFICIENT_EVIDENCE (FR-06)."""
        mock_response = MagicMock()
        mock_response.content = '{"classification": "present", "evidence_citation": null, "confidence_note": "Looks good"}'

        with patch("langchain_groq.ChatGroq") as MockChatGroq:
            mock_llm = MagicMock()
            mock_llm.invoke.return_value = mock_response
            MockChatGroq.return_value = mock_llm

            with patch("app.services.analyzer.get_settings") as mock_settings:
                mock_settings.return_value.groq_api_key = "test-key"
                mock_settings.return_value.llm_model = "llama-3.3-70b-versatile"

                chunks = [ResumeChunk(chunk_id="s:0", source_section=Section.SKILLS, raw_text="Python")]
                reqs = [JDRequirement(requirement_id="r1", requirement_text="Python")]
                retrieval = [
                    RetrievalResult(
                        requirement_id="r1",
                        retrieved_chunk_ids=["s:0"],
                        similarity_scores=[0.80],
                        threshold_pass=True,
                    )
                ]
                results = classify_requirements(reqs, retrieval, chunks)

                assert results[0].classification == Classification.INSUFFICIENT_EVIDENCE

    def test_fallback_to_rule_based_on_llm_error(self):
        """LLM exception should fall back to rule-based classification."""
        with patch("langchain_groq.ChatGroq") as MockChatGroq:
            mock_llm = MagicMock()
            mock_llm.invoke.side_effect = Exception("API error")
            MockChatGroq.return_value = mock_llm

            with patch("app.services.analyzer.get_settings") as mock_settings:
                mock_settings.return_value.groq_api_key = "test-key"
                mock_settings.return_value.llm_model = "llama-3.3-70b-versatile"

                chunks = [ResumeChunk(chunk_id="s:0", source_section=Section.SKILLS, raw_text="Python")]
                reqs = [JDRequirement(requirement_id="r1", requirement_text="Python")]
                retrieval = [
                    RetrievalResult(
                        requirement_id="r1",
                        retrieved_chunk_ids=["s:0"],
                        similarity_scores=[0.85],
                        threshold_pass=True,
                    )
                ]
                results = classify_requirements(reqs, retrieval, chunks)

                # Should fall back to rule-based PRESENT
                assert results[0].classification == Classification.PRESENT

    def test_no_api_key_uses_rule_based(self):
        """Without GROQ_API_KEY, should use rule-based classification."""
        with patch("app.services.analyzer.get_settings") as mock_settings:
            mock_settings.return_value.groq_api_key = ""

            chunks = [ResumeChunk(chunk_id="s:0", source_section=Section.SKILLS, raw_text="Python")]
            reqs = [JDRequirement(requirement_id="r1", requirement_text="Python")]
            retrieval = [
                RetrievalResult(
                    requirement_id="r1",
                    retrieved_chunk_ids=["s:0"],
                    similarity_scores=[0.85],
                    threshold_pass=True,
                )
            ]
            results = classify_requirements(reqs, retrieval, chunks)

            assert results[0].classification == Classification.PRESENT


# ---------------------------------------------------------------------------
# End-to-end pipeline test (Day 4 wired)
# ---------------------------------------------------------------------------


class TestPipelineDay4:
    """End-to-end pipeline with real embeddings, retrieval, and classification."""

    def test_full_pipeline_returns_analyses(self):
        """Pipeline should return analyses with classifications."""
        from io import BytesIO

        from docx import Document

        from app.services.pipeline import run_analysis

        # Build a minimal DOCX resume
        doc = Document()
        for line in [
            "John Doe",
            "SKILLS",
            "Python, FastAPI, PostgreSQL, Docker",
            "EXPERIENCE",
            "Backend Engineer at Acme Corp — built REST APIs with FastAPI",
            "PROJECTS",
            "Resume Analyzer — RAG app comparing resumes to JDs",
            "EDUCATION",
            "B.Tech CSE, GITAM University, 2023-2027",
        ]:
            doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)

        jd = "Requirements: Python (must have), FastAPI experience, Docker containerization"
        report = run_analysis(buf.getvalue(), "resume.docx", jd)

        # Should have requirements, retrieval, and analyses
        assert len(report.requirements) > 0
        assert len(report.retrieval) > 0
        assert len(report.analyses) > 0

        # Every analysis should have a valid classification
        for a in report.analyses:
            assert a.classification in Classification

        # Present/Partial claims must have evidence citations (FR-06)
        for a in report.analyses:
            if a.classification in (Classification.PRESENT, Classification.PARTIAL):
                assert a.evidence_citation is not None, (
                    f"FR-06 violation: {a.classification.value} without citation"
                )

    def test_missing_requirements_get_missing_classification(self):
        """Requirements with no resume coverage should be MISSING."""
        from io import BytesIO

        from docx import Document

        from app.services.pipeline import run_analysis

        doc = Document()
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Python, Flask")
        buf = BytesIO()
        doc.save(buf)

        jd = "Requirements: Kubernetes, Terraform, AWS Lambda"
        report = run_analysis(buf.getvalue(), "resume.docx", jd)

        # All should be MISSING since resume has no cloud/infra content
        assert len(report.analyses) > 0
        assert all(
            a.classification in (Classification.MISSING, Classification.INSUFFICIENT_EVIDENCE)
            for a in report.analyses
        )
